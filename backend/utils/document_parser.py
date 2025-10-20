"""
Utility for extracting plain text from uploaded documents.
"""

from io import BytesIO
import logging
from pathlib import Path
from typing import Callable, Dict

from PyPDF2 import PdfReader
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


class DocumentParser:
    """
    Extracts plain text from supported document formats.
    """

    _TEXT_EXTENSIONS = {".txt", ".md", ".csv", ".tsv"}
    _PDF_EXTENSIONS = {".pdf"}
    _DOCX_EXTENSIONS = {".docx"}

    def __init__(self) -> None:
        self._parsers: Dict[str, Callable[[bytes], str]] = {}

        for ext in self._TEXT_EXTENSIONS:
            self._parsers[ext] = self._parse_text

        for ext in self._PDF_EXTENSIONS:
            self._parsers[ext] = self._parse_pdf

        for ext in self._DOCX_EXTENSIONS:
            self._parsers[ext] = self._parse_docx

    def parse(self, filename: str, content: bytes) -> str:
        """
        Convert uploaded file bytes to plain text.

        Args:
            filename: Name of the uploaded file (used for extension detection).
            content: Raw file bytes.

        Returns:
            Extracted plain text.

        Raises:
            ValueError: If the file type is unsupported or parsing fails.
        """
        suffix = Path(filename).suffix.lower()

        if suffix not in self._parsers:
            raise ValueError(f"Unsupported file type: {suffix or 'unknown'}")

        parser = self._parsers[suffix]
        logger.info("Parsing document '%s' as %s", filename, suffix)
        text = parser(content)

        if not text.strip():
            logger.warning("Parsed document '%s' but extracted text is empty", filename)

        return text

    def _parse_text(self, content: bytes) -> str:
        """Decode plain text-like content using UTF-8 with fallback."""
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError:
            logger.warning("UTF-8 decode failed, falling back to latin-1")
            return content.decode("latin-1", errors="ignore")

    def _parse_pdf(self, content: bytes) -> str:
        """Extract text from a PDF document using PyPDF2."""
        reader = PdfReader(BytesIO(content))

        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception as exc:
                raise ValueError("Unable to decrypt encrypted PDF") from exc

        pages = []
        for index, page in enumerate(reader.pages, start=1):
            try:
                page_text = page.extract_text() or ""
            except Exception as exc:  # PyPDF2 can raise generic exceptions
                logger.error("Failed to extract text from PDF page %s: %s", index, exc)
                page_text = ""

            if page_text:
                pages.append(page_text.strip())

        return "\n\n".join(pages)

    def _parse_docx(self, content: bytes) -> str:
        """Extract text from a DOCX document using python-docx."""
        document = DocxDocument(BytesIO(content))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

        # Tables can contain important text; include them
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" ".join(cells))

        return "\n\n".join(paragraphs)
